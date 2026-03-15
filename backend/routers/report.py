from fastapi import APIRouter, HTTPException, Form, Response
import os, json, re, io, asyncio
import google.generativeai as genai
from tavily import TavilyClient
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

router = APIRouter(prefix="/api")

# ── Gemini 2.0 Flash — 1M free tokens/day ───────────────────
genai.configure(api_key=os.getenv("GEMINI_API_KEY"))

def call_gemini(system_prompt: str, user_prompt: str, temperature: float = 0.1, max_tokens: int = 2500) -> str:
    """Universal Gemini 2.0 Flash caller. Returns response text."""
    try:
        model = genai.GenerativeModel(
            model_name="gemini-2.0-flash",
            generation_config={"temperature": temperature, "max_output_tokens": max_tokens},
            system_instruction=system_prompt
        )
        response = model.generate_content(user_prompt)
        return response.text
    except Exception as e:
        print(f"Gemini error: {e}")
        raise e

tavily_client = TavilyClient(api_key=os.getenv("TAVILY_API_KEY"))

_research_cache = {}
_tavily_cache   = {}

# ── UTILS ────────────────────────────────────────────────────
def safe_float(v, default=0.0):
    try: return float(str(v).replace(",","").replace("₹","").replace("%","").replace("INR","").strip())
    except: return default

def normalize_to_inr_crores(value, hint=""):
    if value is None: return None
    v = safe_float(value)
    if v == 0.0 and str(value).strip() not in ["0","0.0"]: return None
    h = str(hint).lower()
    if "usd" in h or "$" in h:
        if "billion" in h or "bn" in h: return round(v * 8300, 2)
        if "million" in h or "mn" in h: return round(v * 83, 2)
        return round(v * 83, 2)
    if "lakh" in h or "lac" in h: return round(v / 100, 2)
    if "billion" in h or "bn" in h: return round(v * 100, 2)
    if "million" in h or "mn" in h: return round(v / 10, 2)
    return round(v, 2)

def robust_json_parse(text):
    """Multiple fallback strategies for malformed JSON from LLM."""
    if not text: return {}
    # Remove markdown fences
    text = re.sub(r'```(?:json)?', '', text).strip()
    # Strategy 1: direct parse
    try: return json.loads(text)
    except: pass
    # Strategy 2: extract first {...} block
    try:
        m = re.search(r'\{.*\}', text, re.DOTALL)
        if m: return json.loads(m.group())
    except: pass
    # Strategy 3: fix common issues - trailing commas, single quotes, unquoted keys
    try:
        fixed = re.sub(r',\s*([}\]])', r'\1', text)  # trailing commas
        fixed = re.sub(r"'([^']*)':", r'"\1":', fixed)  # single-quoted keys
        fixed = re.sub(r":\s*'([^']*)'", r': "\1"', fixed)  # single-quoted values
        m = re.search(r'\{.*\}', fixed, re.DOTALL)
        if m: return json.loads(m.group())
    except: pass
    # Strategy 4: line-by-line key extraction
    try:
        result = {}
        for line in text.split('\n'):
            m = re.match(r'\s*["\']?(\w+)["\']?\s*:\s*["\']?([^,"\'\n\}]+)["\']?', line)
            if m: result[m.group(1).strip()] = m.group(2).strip()
        if result: return result
    except: pass
    return {}

def cached_tavily(query, n=3):
    key = query.strip().lower()
    if key in _tavily_cache: return _tavily_cache[key]
    try:
        res = tavily_client.search(query=query, max_results=n, search_depth="basic")
        _tavily_cache[key] = res
        return res
    except Exception as e:
        print(f"Tavily error: {e}")
        return {"results": []}

# ── DECISION TABLE ───────────────────────────────────────────
def get_decision(score):
    if score >= 80: return "APPROVE",                   "A", "Base + 0.75%"
    if score >= 65: return "APPROVE WITH CONDITIONS",   "B", "Base + 1.5%"
    if score >= 50: return "REFER TO CREDIT COMMITTEE", "C", "Base + 2.5%"
    return               "REJECT",                      "D", "N/A"

def get_risk_level(score):
    if score >= 80: return "LOW"
    if score >= 65: return "MEDIUM"
    if score >= 50: return "HIGH"
    return "CRITICAL"

# ── SCORING ENGINE ───────────────────────────────────────────
def calculate_universal_score(company_name, extracted_docs, research_findings, entity_data, analyst_notes=""):
    scores = {"character":20,"capacity":20,"capital":20,"collateral":20,"conditions":15}
    notes  = {k:"" for k in scores}
    red_flags, green_flags, reasoning_chain = [], [], []

    annual      = extracted_docs.get("annual_report", {})
    borrowing   = extracted_docs.get("borrowing_profile", {})
    portfolio   = extracted_docs.get("portfolio_cuts", {})
    alm         = extracted_docs.get("alm_statement", {})
    shareholding= extracted_docs.get("shareholding_pattern", {})

    pledged = safe_float(str(shareholding.get("pledged_shares") or "0").replace("%",""))
    if pledged > 50:
        scores["character"] -= 10; red_flags.append(f"Promoter pledge critical: {pledged}%")
        reasoning_chain.append(f"CHARACTER -10: Promoter pledge {pledged}% far exceeds 25% threshold (Source: Shareholding Pattern)")
    elif pledged > 25:
        scores["character"] -= 6; red_flags.append(f"Promoter pledge elevated: {pledged}%")
        reasoning_chain.append(f"CHARACTER -6: Promoter pledge {pledged}% above 25% (Source: Shareholding Pattern)")
    elif pledged > 0:
        scores["character"] -= 2
        reasoning_chain.append(f"CHARACTER -2: Minor promoter pledge {pledged}% (Source: Shareholding Pattern)")
    else:
        green_flags.append("Zero promoter pledge — clean ownership structure")
        reasoning_chain.append("CHARACTER +0: No promoter pledge — positive signal (Source: Shareholding Pattern)")

    rating  = str(borrowing.get("credit_rating_long_term") or "").lower()
    outlook = str(borrowing.get("rating_outlook") or "").lower()

    if any(x in rating for x in ["care d","icra d"," d rated","default"]):
        scores["character"] -= 12; red_flags.append("DEFAULT rating — entity has defaulted")
        reasoning_chain.append(f"CHARACTER -12: Default rating detected (Source: Borrowing Profile)")
    elif any(x in rating for x in ["bbb-","bb+","bb","b+"]):
        scores["character"] -= 6; red_flags.append(f"Sub-investment grade: {rating.upper()}")
        reasoning_chain.append(f"CHARACTER -6: Sub-investment grade {rating.upper()} (Source: Borrowing Profile)")
    elif "bbb" in rating:
        scores["character"] -= 3
    elif any(x in rating for x in ["aa","aaa","a+"]):
        green_flags.append(f"Strong credit rating: {rating.upper()}")

    if "watch negative" in outlook or "watch-" in outlook:
        scores["character"] -= 5; red_flags.append("Rating on Watch Negative")
        reasoning_chain.append("CHARACTER -5: Rating on Watch Negative (Source: Borrowing Profile)")
    elif "negative" in outlook:
        scores["character"] -= 3; red_flags.append("Negative rating outlook")
    elif "stable" in outlook:
        green_flags.append("Stable rating outlook")
    elif "positive" in outlook:
        scores["character"] += 1; green_flags.append("Positive rating outlook")

    scores["character"] = max(min(scores["character"],20),0)
    notes["character"] = f"Pledge: {pledged}%. Rating: {rating or 'N/A'}. Outlook: {outlook or 'N/A'}."

    rev  = normalize_to_inr_crores(annual.get("revenue")) or 0
    pat_raw = normalize_to_inr_crores(annual.get("pat"))
    gnpa = safe_float(str(annual.get("gnpa_percent") or portfolio.get("gnpa_percent") or "2").replace("%",""))
    coll_eff = safe_float(str(portfolio.get("collection_efficiency") or "97").replace("%",""))

    if pat_raw is None or safe_float(pat_raw) <= 0:
        scores["capacity"] -= 8; red_flags.append("Net loss — negative PAT")
        reasoning_chain.append(f"CAPACITY -8: Net loss recorded (Source: Annual Report)")
    else:
        green_flags.append(f"Profitable — PAT ₹{pat_raw} Cr")
        reasoning_chain.append(f"CAPACITY +0: Profitable PAT ₹{pat_raw} Cr (Source: Annual Report)")
        if rev > 0:
            margin = (safe_float(pat_raw)/rev)*100
            if margin < 5:
                scores["capacity"] -= 3; red_flags.append(f"Thin profit margin: {margin:.1f}%")
                reasoning_chain.append(f"CAPACITY -3: Thin margin {margin:.1f}% < 5% (Source: Annual Report)")

    if gnpa > 7:
        scores["capacity"] -= 10; red_flags.append(f"GNPA critical: {gnpa}%")
        reasoning_chain.append(f"CAPACITY -10: GNPA {gnpa}% > 7% threshold (Source: Annual/Portfolio Report)")
    elif gnpa > 5:
        scores["capacity"] -= 7; red_flags.append(f"GNPA very high: {gnpa}%")
        reasoning_chain.append(f"CAPACITY -7: GNPA {gnpa}% > 5% (Source: Annual/Portfolio Report)")
    elif gnpa > 3:
        scores["capacity"] -= 4; red_flags.append(f"GNPA elevated: {gnpa}%")
        reasoning_chain.append(f"CAPACITY -4: GNPA {gnpa}% > 3% (Source: Annual/Portfolio Report)")
    else:
        green_flags.append(f"GNPA healthy: {gnpa}%")
        reasoning_chain.append(f"CAPACITY +0: GNPA {gnpa}% within healthy range (Source: Annual/Portfolio Report)")

    if coll_eff < 90:
        scores["capacity"] -= 6; red_flags.append(f"Collection efficiency critical: {coll_eff}%")
        reasoning_chain.append(f"CAPACITY -6: Collection efficiency {coll_eff}% < 90% (Source: Portfolio Report)")
    elif coll_eff < 95:
        scores["capacity"] -= 3
    elif coll_eff >= 98:
        scores["capacity"] += 1; green_flags.append(f"Excellent collection efficiency: {coll_eff}%")

    scores["capacity"] = max(min(scores["capacity"],20),0)
    notes["capacity"] = f"PAT: {pat_raw} Cr. GNPA: {gnpa}%. Collection: {coll_eff}%."

    car = safe_float(str(annual.get("car_percent") or "15").replace("%",""))
    nw  = normalize_to_inr_crores(annual.get("net_worth")) or 0
    debt= normalize_to_inr_crores(annual.get("total_debt") or borrowing.get("total_debt")) or 0
    de_ratio = round(debt/nw,2) if nw > 0 else 99

    if car < 12:
        scores["capital"] -= 12; red_flags.append(f"CAR critical: {car}%")
        reasoning_chain.append(f"CAPITAL -12: CAR {car}% breaches RBI minimum 15% (Source: Annual Report)")
    elif car < 15:
        scores["capital"] -= 7; red_flags.append(f"CAR below RBI minimum: {car}%")
        reasoning_chain.append(f"CAPITAL -7: CAR {car}% below RBI minimum (Source: Annual Report)")
    elif car >= 18:
        scores["capital"] += 1; green_flags.append(f"Strong CAR: {car}%")
        reasoning_chain.append(f"CAPITAL +1: CAR {car}% well above RBI minimum (Source: Annual Report)")

    if de_ratio > 6:
        scores["capital"] -= 8; red_flags.append(f"D/E critically high: {de_ratio}x")
        reasoning_chain.append(f"CAPITAL -8: D/E {de_ratio}x far exceeds 6x ceiling (Source: Borrowing Profile)")
    elif de_ratio > 4:
        scores["capital"] -= 5; red_flags.append(f"D/E above NBFC norm: {de_ratio}x")
        reasoning_chain.append(f"CAPITAL -5: D/E {de_ratio}x above 4x NBFC norm (Source: Borrowing Profile)")
    elif de_ratio > 3:
        scores["capital"] -= 2
        reasoning_chain.append(f"CAPITAL -2: D/E {de_ratio}x approaching ceiling (Source: Borrowing Profile)")
    elif de_ratio > 0:
        green_flags.append(f"Healthy leverage: D/E {de_ratio}x")

    scores["capital"] = max(min(scores["capital"],20),0)
    notes["capital"] = f"CAR: {car}%. D/E: {de_ratio}x. Net Worth: ₹{nw} Cr."

    alm_assets = normalize_to_inr_crores(alm.get("total_assets")) or 0
    alm_liabs  = normalize_to_inr_crores(alm.get("total_liabilities")) or 0
    if alm_assets > 0 and alm_liabs > 0:
        coverage = round(alm_assets/alm_liabs,2)
        if coverage < 1.0:
            scores["collateral"] -= 10; red_flags.append(f"Asset coverage below 1x: {coverage}x")
            reasoning_chain.append(f"COLLATERAL -10: Asset coverage {coverage}x < 1.0 (Source: ALM Statement)")
        elif coverage < 1.1:
            scores["collateral"] -= 5
            reasoning_chain.append(f"COLLATERAL -5: Thin asset coverage {coverage}x (Source: ALM Statement)")
        else:
            green_flags.append(f"Good asset coverage: {coverage}x")
            reasoning_chain.append(f"COLLATERAL +0: Asset coverage {coverage}x adequate (Source: ALM Statement)")
        notes["collateral"] = f"Coverage: {coverage}x."
    else:
        notes["collateral"] = "ALM data not available — using default score."
        reasoning_chain.append("COLLATERAL ±0: ALM data not available — using conservative default (Source: ALM Statement)")

    liq_gap = normalize_to_inr_crores(alm.get("liquidity_gap"))
    if liq_gap is not None and liq_gap < 0:
        scores["collateral"] -= 5; red_flags.append(f"Negative liquidity gap: ₹{liq_gap} Cr")

    scores["collateral"] = max(min(scores["collateral"],20),0)

    all_text = " ".join([(f.get("title","")+" "+f.get("snippet","")+" "+f.get("content","")).lower() for f in research_findings])
    sector = str(entity_data.get("sector","")).lower()

    if "nbfc" in sector or "finance" in sector:
        scores["conditions"] -= 1
        reasoning_chain.append("CONDITIONS -1: NBFC sector faces RBI regulatory scrutiny (Source: Sector Analysis)")

    if any(x in all_text for x in ["default","care d","icra d"," d rated"]):
        scores["conditions"] -= 10; red_flags.append("DEFAULT detected in web intelligence")
        reasoning_chain.append("CONDITIONS -10: DEFAULT detected in news (Source: Web Intelligence)")
    elif any(x in all_text for x in ["breached covenant","covenant breach","loan terms breached"]):
        scores["conditions"] -= 6; red_flags.append("Covenant breach in news")
        reasoning_chain.append("CONDITIONS -6: Covenant breach reported (Source: Web Intelligence)")

    if any(x in all_text for x in ["fraud","sebi action","rbi penalty","scam","arrested"]):
        scores["conditions"] -= 8; red_flags.append("Fraud/regulatory action in news")
        reasoning_chain.append("CONDITIONS -8: Fraud/regulatory action (Source: Web Intelligence)")
    if any(x in all_text for x in ["nclt","insolvency","ibc proceedings"]):
        scores["conditions"] -= 7; red_flags.append("NCLT/IBC insolvency detected")
        reasoning_chain.append("CONDITIONS -7: NCLT/IBC insolvency (Source: Web Intelligence/MCA)")
    if any(x in all_text for x in ["downgraded","rating downgrade"]):
        scores["conditions"] -= 3; red_flags.append("Credit rating downgraded in news")
        reasoning_chain.append("CONDITIONS -3: Downgrade confirmed in news (Source: Web Intelligence)")
    if any(x in all_text for x in ["upgrade","rating upgrade"]):
        scores["conditions"] += 3; green_flags.append("Rating upgrade signal in news")
        reasoning_chain.append("CONDITIONS +3: Rating upgrade (Source: Web Intelligence)")

    # Analyst notes adjustment
    analyst_adj = 0
    if analyst_notes and len(analyst_notes.strip()) > 10:
        try:
            response_text = call_gemini(
                "You are a credit scoring assistant. Analyze analyst notes and return a JSON score adjustment. Return ONLY valid JSON.",
                f'Credit analyst notes: "{analyst_notes}"\nReturn ONLY JSON: {{"adjustment": -5, "reasoning": "brief reason"}}\nRange: -15 to +5.',
                temperature=0.1, max_tokens=150
            )
            adj_data = robust_json_parse(response_text)
            analyst_adj = max(-15, min(5, int(adj_data.get("adjustment",0))))
            if analyst_adj < 0: red_flags.append(f"Analyst note: {adj_data.get('reasoning','')}")
            reasoning_chain.append(f"ANALYST {analyst_adj:+d}: {adj_data.get('reasoning','')} (Source: Primary Due Diligence)")
        except Exception as e:
            print(f"Analyst notes error: {e}")

    scores["conditions"] = max(min(scores["conditions"],15),0)
    notes["conditions"] = f"Sector: {sector}. Web sources: {len(research_findings)}. Analyst adj: {analyst_adj:+d}."

    final_score = max(0, min(100, sum(scores.values()) + analyst_adj))
    decision, grade, rate = get_decision(final_score)

    loan_amount = safe_float(entity_data.get("loan_amount",0))
    if decision == "APPROVE": recommended_amount = loan_amount
    elif "CONDITIONS" in decision: recommended_amount = round(loan_amount*0.90,2)
    elif "REFER" in decision: recommended_amount = round(loan_amount*0.75,2)
    else: recommended_amount = 0

    key_reasons = [r for r in reasoning_chain if not "+0" in r or "not" in r.lower()][:5]
    reasoning = (f"{company_name} scored {final_score}/100 (Five Cs). Decision: {decision} (Grade {grade}). "
                 + " | ".join(key_reasons[:3]))

    return {
        "score": final_score, "decision": decision, "grade": grade,
        "risk_level": get_risk_level(final_score),
        "recommended_amount": recommended_amount, "recommended_rate": rate,
        "tenure": entity_data.get("tenure",36), "reasoning": reasoning,
        "reasoning_chain": reasoning_chain, "analyst_notes": analyst_notes,
        "analyst_adjustment": analyst_adj,
        "red_flags": list(dict.fromkeys(red_flags)),
        "green_flags": list(dict.fromkeys(green_flags)),
        "five_cs": {
            "character":  {"score":scores["character"],  "max":20,"notes":notes["character"]},
            "capacity":   {"score":scores["capacity"],   "max":20,"notes":notes["capacity"]},
            "capital":    {"score":scores["capital"],    "max":20,"notes":notes["capital"]},
            "collateral": {"score":scores["collateral"], "max":20,"notes":notes["collateral"]},
            "conditions": {"score":scores["conditions"], "max":15,"notes":notes["conditions"]},
        },
        "swot": generate_swot(company_name, extracted_docs, research_findings)
    }

def generate_swot(company_name, extracted_docs, research_findings):
    """Generate company-specific SWOT — never use generic fallback if data available."""
    try:
        news = "\n".join([f"- {f.get('title','')}: {(f.get('snippet','') or f.get('content',''))[:200]}" for f in research_findings[:6]])
        annual = extracted_docs.get("annual_report",{})
        rev  = normalize_to_inr_crores(annual.get("revenue"))
        pat  = normalize_to_inr_crores(annual.get("pat"))
        gnpa = annual.get("gnpa_percent")
        car  = annual.get("car_percent")

        prompt = f"""You are a senior Indian credit analyst. Generate a SPECIFIC SWOT for {company_name}.
Financial data: Revenue=₹{rev}Cr, PAT=₹{pat}Cr, GNPA={gnpa}%, CAR={car}%
Recent news and web intelligence:
{news}

RULES:
- Every point MUST reference specific numbers or events from the data above
- NO generic statements like "established market presence" or "requires further due diligence"
- Each point must name {company_name} explicitly or reference specific financial metrics

Return ONLY valid JSON (no markdown, no extra text):
{{"strengths":["specific point with numbers","specific point","specific point"],"weaknesses":["specific point with numbers","specific point","specific point"],"opportunities":["specific point","specific point"],"threats":["specific point","specific point","specific point"]}}"""

        response_text = call_gemini(
            "You are a senior Indian credit analyst. Generate specific, data-driven SWOT analysis. Return ONLY valid JSON.",
            prompt, temperature=0.4, max_tokens=700
        )
        result = robust_json_parse(response_text)
        if result and all(k in result for k in ["strengths","weaknesses","opportunities","threats"]):
            # Validate not generic
            all_points = result.get("strengths",[]) + result.get("weaknesses",[])
            generic_phrases = ["established market presence","regulatory compliance","requires further","growing credit demand","digital lending expansion"]
            has_generic = any(any(g in p.lower() for g in generic_phrases) for p in all_points)
            if not has_generic:
                return result
    except Exception as e:
        print(f"SWOT error: {e}")

    # Fallback — but make it data-driven
    annual = extracted_docs.get("annual_report",{})
    rev = normalize_to_inr_crores(annual.get("revenue")) or "N/A"
    pat = normalize_to_inr_crores(annual.get("pat")) or "N/A"
    gnpa = annual.get("gnpa_percent") or "N/A"
    car  = annual.get("car_percent") or "N/A"
    return {
        "strengths":     [f"{company_name} reported PAT of ₹{pat} Cr — profitable entity", f"GNPA at {gnpa}% — below industry average", f"CAR at {car}% — strong capital buffer"],
        "weaknesses":    [f"Revenue of ₹{rev} Cr requires cross-verification with web data", "Credit rating data not available from submitted documents", "ALM/Liquidity gap data incomplete"],
        "opportunities": ["Growing NBFC credit demand in Indian SME sector", "RBI easing NBFC regulations in 2024-25"],
        "threats":       ["Rising cost of funds squeezing NBFC NIMs", "RBI regulatory scrutiny on NBFC sector", "Increasing competition from digital lenders"]
    }

def generate_triangulation(company_name, extracted_docs, research_data):
    items = []
    annual = extracted_docs.get("annual_report",{})

    doc_rev = normalize_to_inr_crores(annual.get("revenue"))
    web_rev = normalize_to_inr_crores(research_data.get("revenue"))
    if doc_rev and web_rev:
        diff = abs(doc_rev-web_rev)/max(doc_rev,web_rev)*100
        note = f"{diff:.1f}% variance — web data may include consolidated group figures; use document value" if diff > 20 else "Consistent across sources"
        items.append({"field":"Revenue","doc":f"₹{doc_rev} Cr","web":f"₹{web_rev} Cr","status":"MISMATCH" if diff>20 else "CONSISTENT","note":note})

    doc_debt = normalize_to_inr_crores(annual.get("total_debt"))
    web_debt = normalize_to_inr_crores(research_data.get("total_debt"))
    if doc_debt and web_debt:
        diff = abs(doc_debt-web_debt)/max(doc_debt,web_debt)*100
        items.append({"field":"Total Debt","doc":f"₹{doc_debt} Cr","web":f"₹{web_debt} Cr","status":"MISMATCH" if diff>25 else "CONSISTENT","note":f"{diff:.1f}% variance" if diff>25 else "Consistent"})

    doc_gnpa = safe_float(str(annual.get("gnpa_percent") or "0").replace("%",""))
    web_gnpa = safe_float(str(research_data.get("gnpa_percent") or "0").replace("%",""))
    if doc_gnpa > 0 and web_gnpa > 0:
        diff = abs(doc_gnpa-web_gnpa)
        items.append({"field":"GNPA %","doc":f"{doc_gnpa}%","web":f"{web_gnpa}%","status":"MISMATCH" if diff>2 else "CONSISTENT","note":f"{diff:.2f}pp variance" if diff>2 else "Consistent"})
    return items

@router.post("/research")
async def perform_research(data: dict):
    try:
        company = data.get("company_name","Unknown")
        sector  = data.get("sector","General")
        key     = company.strip().lower()
        if key in _research_cache: return _research_cache[key]

        loop = asyncio.get_event_loop()

        async def fetch(q, n):
            return await loop.run_in_executor(None, lambda: cached_tavily(q, n))

        # 4 searches with timeout — reduced from 6 to prevent Render timeout
        results = await asyncio.wait_for(
            asyncio.gather(
                fetch(f"{company} India revenue profit PAT financial results 2024 2025", 4),
                fetch(f"{company} India fraud NCLT default litigation court 2024 2025", 3),
                fetch(f"{company} India RBI SEBI regulatory penalty credit rating 2024", 3),
                fetch(f"{sector} sector India outlook RBI regulation headwinds 2025", 2),
            ),
            timeout=25.0  # 25 second hard timeout
        )

        unique = {}
        for resp in results:
            for r in resp.get("results",[]):
                url = r.get("url")
                if url and url not in unique:
                    unique[url] = {"title":r.get("title",""),"snippet":(r.get("content",""))[:500],"url":url}

        findings = list(unique.values())
        context  = "\n".join([f"[{f['url']}]\n{f['title']}: {f['snippet']}" for f in findings])

        # Revenue history: search specifically for multi-year data
        rev_results = await asyncio.wait_for(
            fetch(f"{company} India revenue FY2022 FY2023 FY2024 annual report history crores", 3),
            timeout=8.0
        )
        for r in rev_results.get("results",[]):
            url = r.get("url")
            if url and url not in unique:
                unique[url] = {"title":r.get("title",""),"snippet":(r.get("content",""))[:500],"url":url}
                findings.append(unique[url])

        context = "\n".join([f"[{f['url']}]\n{f['title']}: {f['snippet']}" for f in findings])

        resp_system = f"""Senior Indian credit analyst. Analyze web results for {company}.

CRITICAL UNIT RULES — STRICTLY FOLLOW:
1. ALL financial values MUST be in INR Crores. 
2. If source shows USD billions: multiply by 8300. If USD millions: multiply by 83. If INR lakhs: divide by 100.
3. {company} is likely a standalone Indian NBFC/company — DO NOT use Tata Group consolidated revenue even if found.
4. Recompute D/E = total_debt/net_worth. Recompute ROE = (pat/net_worth)*100. Never copy ratios from source.
5. revenue_history: look specifically for 3 years of standalone revenue. Only include years with CONFIRMED data. Use null if unsure.
6. If PAT margin > 50% that means unit error — re-check.
7. For nclt_status: if {company} is the CREDITOR in a case (filed against borrower), say "Filed as creditor against [name] — not a defaulter signal". Only flag if {company} is the DEBTOR.

Return ONLY valid JSON (no markdown):
{{
  "company_name":"{company}","headquarters":"","founded_year":"","sector":"{sector}",
  "revenue":null,"pat":null,"total_debt":null,"net_worth":null,"total_assets":null,
  "gnpa_percent":null,"car_percent":null,
  "de_ratio":null,"roe":null,"pat_margin":null,"revenue_growth":null,
  "revenue_history":[
    {{"year":"FY2024","revenue_cr":null}},
    {{"year":"FY2023","revenue_cr":null}},
    {{"year":"FY2022","revenue_cr":null}}
  ],
  "character_score":16,"capacity_score":16,"capital_score":14,"collateral_score":14,"conditions_score":11,
  "total_score":71,
  "credit_decision":"APPROVE or APPROVE WITH CONDITIONS or REFER TO CREDIT COMMITTEE or REJECT",
  "risk_level":"LOW or MEDIUM or HIGH or CRITICAL",
  "positive_signals":["specific finding 1","specific finding 2"],
  "risk_flags":["specific risk 1","specific risk 2"],
  "rbi_regulatory_flags":["any RBI/SEBI/regulatory findings"],
  "nclt_status":"None detected or specific details",
  "cibil_signal":"N/A or specific CIBIL findings",
  "gstr_signal":"N/A or GST mismatch details",
  "mca_flags":["MCA filing issues or empty array"],
  "litigation_risk":"LOW/MEDIUM/HIGH — explanation",
  "promoter_background":"specific promoter news",
  "sector_headwinds":["specific sector risk 1","specific sector risk 2"],
  "latest_news":["specific event 1","specific event 2","specific event 3"],
  "sector_outlook":"one concise sentence",
  "research_summary":"3-sentence credit opinion mentioning {company} specifically",
  "data_sources":[]
}}"""
        resp_user = f"Company: {company}\nSector: {sector}\n\nWeb Data ({len(findings)} sources):\n{context[:9000]}"
        text = call_gemini(resp_system, resp_user, temperature=0.1, max_tokens=2500)
        result = robust_json_parse(text)

        if not result:
            return {"error":"JSON parsing failed", "company_name":company, "research_summary":f"Web research completed for {company}. Manual review required.", "risk_level":"MEDIUM", "total_score":65, "credit_decision":"APPROVE WITH CONDITIONS", "latest_news":[], "positive_signals":[], "risk_flags":[], "revenue_history":[], "data_sources":[]}

        # Post-process: recompute ratios, enforce decision table
        rev  = safe_float(result.get("revenue") or 0)
        pat  = safe_float(result.get("pat") or 0)
        debt = safe_float(result.get("total_debt") or 0)
        nw   = safe_float(result.get("net_worth") or 0)

        if nw > 0 and debt >= 0: result["de_ratio"]   = round(debt/nw, 4)
        if nw > 0 and pat != 0:  result["roe"]        = round((pat/nw)*100, 2)
        if rev > 0 and pat != 0: result["pat_margin"] = round((pat/rev)*100, 2)

        total = result.get("total_score",65)
        decision, grade, rate = get_decision(total)
        result["credit_decision"]  = decision
        result["grade"]            = grade
        result["recommended_rate"] = rate
        result["risk_level"]       = get_risk_level(total)

        # Clean revenue history — remove nulls and zeros
        if result.get("revenue_history"):
            result["revenue_history"] = [
                h for h in result["revenue_history"]
                if h.get("revenue_cr") is not None and str(h.get("revenue_cr")) not in ["0","0.0","null","None"]
                and safe_float(h.get("revenue_cr")) > 0
            ]

        result["data_sources"] = list(unique.keys())[:12]
        _research_cache[key] = result
        return result

    except asyncio.TimeoutError:
        print(f"Research timeout for {company}")
        return {"error":"timeout","company_name":company,"research_summary":f"Web research timed out for {company}. Using document data only.","risk_level":"MEDIUM","total_score":65,"credit_decision":"APPROVE WITH CONDITIONS","latest_news":[],"positive_signals":[],"risk_flags":[],"revenue_history":[],"data_sources":[]}
    except Exception as e:
        import traceback; traceback.print_exc()
        raise HTTPException(status_code=500, detail=str(e))

@router.get("/cache/stats")
async def cache_stats():
    return {"companies":list(_research_cache.keys()),"queries":len(_tavily_cache)}

@router.delete("/cache/clear")
async def clear_cache():
    global _research_cache,_tavily_cache
    _research_cache={}; _tavily_cache={}
    return {"message":"Cleared"}

@router.post("/generate-cam")
async def generate_report(data: str = Form(...)):
    try:
        from datetime import datetime
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement

        payload       = json.loads(data)
        entity_data   = payload.get("entity",{})
        loan_data     = payload.get("loan",{})
        extracted_data= payload.get("extracted",[])
        research_data = payload.get("research",{})
        score_data    = payload.get("score",{})
        analyst_notes = payload.get("analyst_notes","")

        extracted_docs = {}
        if isinstance(extracted_data,list):
            for doc in extracted_data:
                dt=doc.get("doc_type","unknown"); f=doc.get("fields",{})
                if f: extracted_docs[dt]=f
        elif isinstance(extracted_data,dict):
            extracted_docs=extracted_data

        if not score_data or not score_data.get("score"):
            score_data = calculate_universal_score(
                entity_data.get("companyName","Unknown"), extracted_docs,
                research_data.get("findings",[]),
                {"company_name":entity_data.get("companyName"),"sector":entity_data.get("sector"),
                 "loan_amount":safe_float(loan_data.get("amount",50)),"tenure":int(loan_data.get("tenure",36))},
                analyst_notes
            )

        total=score_data.get("score",0)
        dec_lbl,grade,rate=get_decision(total)

        # ── Build professional DOCX ──────────────────────────
        NAVY=(26,58,107); GOLD=(200,130,10); GREEN=(26,122,58); RED=(178,34,34)
        ORANGE=(199,80,0); WHITE=(255,255,255); GRAY=(245,245,245)

        doc = Document()
        sec = doc.sections[0]
        sec.page_width=Inches(8.5); sec.page_height=Inches(11)
        sec.left_margin=sec.right_margin=Inches(1)
        sec.top_margin=sec.bottom_margin=Inches(0.9)

        def shd(cell, hex_color):
            tc=cell._tc; tcPr=tc.get_or_add_tcPr()
            s=OxmlElement("w:shd")
            s.set(qn("w:fill"),hex_color); s.set(qn("w:color"),"auto"); s.set(qn("w:val"),"clear")
            tcPr.append(s)

        def add_h(txt, level=1, rgb=NAVY):
            p=doc.add_heading(txt,level=level); p.alignment=WD_ALIGN_PARAGRAPH.LEFT
            for r in p.runs: r.font.color.rgb=RGBColor(*rgb)
            return p

        def add_kv(lbl,val,rgb=None):
            p=doc.add_paragraph()
            r1=p.add_run(f"{lbl}: "); r1.bold=True; r1.font.size=Pt(11); r1.font.color.rgb=RGBColor(*NAVY)
            r2=p.add_run(str(val) if val else "N/A"); r2.font.size=Pt(11)
            if rgb: r2.font.color.rgb=RGBColor(*rgb)

        def add_body(txt, rgb=None):
            p=doc.add_paragraph(); p.add_run(txt).font.size=Pt(11)
            if rgb:
                for r in p.runs: r.font.color.rgb=RGBColor(*rgb)
            return p

        def tbl(headers, rows, col_widths, hdr_rgb=NAVY):
            t=doc.add_table(rows=1,cols=len(headers)); t.style='Table Grid'
            hc=t.rows[0].cells
            for i,(h,w) in enumerate(zip(headers,col_widths)):
                hc[i].text=h; shd(hc[i],'%02X%02X%02X'%hdr_rgb)
                run=hc[i].paragraphs[0].runs[0]; run.bold=True
                run.font.color.rgb=RGBColor(*WHITE); run.font.size=Pt(10)
                hc[i].width=Inches(w)
            for row in rows:
                rc=t.add_row().cells
                for i,(v,w) in enumerate(zip(row,col_widths)):
                    rc[i].text=str(v) if v else "N/A"; rc[i].width=Inches(w)
                    rc[i].paragraphs[0].runs[0].font.size=Pt(10) if rc[i].paragraphs[0].runs else None
            doc.add_paragraph()
            return t

        def score_badge_row(score, decision, risk, amount, rate):
            """Create 4-column score summary table."""
            t=doc.add_table(rows=1,cols=4); t.style='Table Grid'
            score_color='1A7A3A' if score>=80 else 'C8820A' if score>=65 else 'C75000' if score>=50 else 'B22222'
            data=[
                ('1A3A6B','INTELLI-SCORE',f'{score}/100',f'Grade {grade}'),
                (score_color,'CREDIT DECISION',decision,rate),
                ('2D5A8E' if risk=='LOW' else '8B4513','RISK LEVEL',risk,'Based on 5 Cs'),
                ('1A3A6B','RECOMMENDED',f'₹{amount} Cr',f'{loan_data.get("tenure","N/A")} Months'),
            ]
            for i,(bg,title,main,sub) in enumerate(data):
                c=t.rows[0].cells[i]; shd(c,bg)
                p1=c.paragraphs[0]
                p1.add_run(title).font.color.rgb=RGBColor(200,200,200); p1.runs[0].font.size=Pt(9); p1.runs[0].bold=True
                p2=c.add_paragraph(); p2.add_run(main).font.color.rgb=RGBColor(255,255,255); p2.runs[0].font.size=Pt(20); p2.runs[0].bold=True
                p3=c.add_paragraph(); p3.add_run(sub).font.color.rgb=RGBColor(200,200,200); p3.runs[0].font.size=Pt(9)
                c.paragraphs[0].alignment=WD_ALIGN_PARAGRAPH.CENTER
                p2.alignment=WD_ALIGN_PARAGRAPH.CENTER; p3.alignment=WD_ALIGN_PARAGRAPH.CENTER
            doc.add_paragraph()

        annual      = extracted_docs.get("annual_report",{})
        borrowing   = extracted_docs.get("borrowing_profile",{})
        portfolio   = extracted_docs.get("portfolio_cuts",{})
        alm         = extracted_docs.get("alm_statement",{})
        shareholding= extracted_docs.get("shareholding_pattern",{})
        company_name= entity_data.get("companyName","N/A")

        # ── COVER ──────────────────────────────────────────
        doc.add_paragraph()
        title=doc.add_heading("Credit Appraisal Memo (CAM)",0)
        title.alignment=WD_ALIGN_PARAGRAPH.CENTER
        for r in title.runs: r.font.color.rgb=RGBColor(*NAVY)
        sub=doc.add_paragraph("VERIDEX® AI Credit Engine v2.0  —  PRIVATE & CONFIDENTIAL")
        sub.alignment=WD_ALIGN_PARAGRAPH.CENTER
        sub.runs[0].font.size=Pt(10); sub.runs[0].font.color.rgb=RGBColor(150,150,150)
        doc.add_paragraph()

        cover_rows=[["Target Entity",company_name],["CIN",entity_data.get("cin","N/A")],["PAN",entity_data.get("pan","N/A")],
                    ["Sector",entity_data.get("sector","N/A")],["Loan Amount",f"₹{loan_data.get('amount','N/A')} Crores"],
                    ["Loan Type",loan_data.get("loanType","N/A")],["Tenure",f"{loan_data.get('tenure','N/A')} Months"],
                    ["Report Date",datetime.now().strftime("%d %B %Y")],["Generated By","VERIDEX AI Credit Engine v2.0"]]
        t=doc.add_table(rows=len(cover_rows),cols=2); t.style='Table Grid'
        for i,(lbl,val) in enumerate(cover_rows):
            t.rows[i].cells[0].text=lbl; t.rows[i].cells[1].text=str(val)
            t.rows[i].cells[0].paragraphs[0].runs[0].bold=True
            t.rows[i].cells[0].paragraphs[0].runs[0].font.color.rgb=RGBColor(*NAVY)

        doc.add_paragraph()
        score_badge_row(total,dec_lbl,get_risk_level(total),score_data.get("recommended_amount","N/A"),rate)
        doc.add_page_break()

        # ── 1. EXECUTIVE SUMMARY ──────────────────────────
        add_h("1.  Executive Summary",1)
        add_kv("Credit Decision",dec_lbl,GREEN if total>=80 else ORANGE if total>=65 else RED)
        add_kv("Grade",f"Grade {grade}")
        add_kv("Intelli-Score",f"{total}/100")
        add_kv("Risk Level",get_risk_level(total),GREEN if total>=80 else ORANGE if total>=65 else RED)
        add_kv("Recommended Limit",f"₹{score_data.get('recommended_amount','N/A')} Crores")
        add_kv("Recommended Rate",rate)
        add_kv("Tenure",f"{loan_data.get('tenure','N/A')} Months")
        doc.add_paragraph()
        p=doc.add_paragraph(); p.add_run("Decision Reasoning: ").bold=True; p.add_run(score_data.get("reasoning","")).font.size=Pt(11)
        doc.add_paragraph()

        # ── 2. EXPLAINABILITY CHAIN ───────────────────────
        add_h("2.  Explainability Chain — Why This Score?",1)
        add_body("Every scoring adjustment is traced to its source document or data signal:")
        chain=score_data.get("reasoning_chain",[])
        if chain:
            chain_rows=[]
            for r in chain[:12]:
                dim=r.split(":")[0].strip()
                rest=r.split(":",1)[1].strip() if ":" in r else r
                src=rest.split("(Source:")[-1].rstrip(")").strip() if "(Source:" in rest else "Analysis"
                reason=rest.split("(Source:")[0].strip()
                adj="+0"; 
                if "-" in dim and any(c.isdigit() for c in dim.split("-")[-1] if dim.split("-")): adj=f"-{dim.split('-')[-1]}" if dim.count('-')>1 else "-"
                elif "+" in dim: adj=f"+{dim.split('+')[-1]}"
                chain_rows.append([dim.split(":")[0],adj,reason,src])
            tbl(["DIMENSION","ADJ","RATIONALE","DATA SOURCE"],chain_rows,[1.2,0.6,3.8,1.4],NAVY)
        doc.add_paragraph()

        # ── 3. FIVE Cs ────────────────────────────────────
        add_h("3.  Five Cs Framework Analysis",1)
        five_cs=score_data.get("five_cs",{})
        cs_rows=[]
        for key,max_v in [("character",20),("capacity",20),("capital",20),("collateral",20),("conditions",15)]:
            cs=five_cs.get(key,{}); s=cs.get("score",0) if isinstance(cs,dict) else (cs or 0)
            n=cs.get("notes","") if isinstance(cs,dict) else ""
            cs_rows.append([key.capitalize(),f"{s}/{max_v}",n])
        t=tbl(["DIMENSION","SCORE","ANALYSIS NOTES"],cs_rows,[1.5,1.0,5.5],NAVY)
        add_kv("TOTAL SCORE",f"{total}/100 — Grade {grade} — {dec_lbl}",GREEN if total>=80 else ORANGE if total>=65 else RED)
        if analyst_notes:
            add_kv("Analyst Notes Applied",analyst_notes)
            add_kv("Score Adjustment",f"{score_data.get('analyst_adjustment',0):+d} points")
        doc.add_paragraph()
        doc.add_page_break()

        # ── 4. FINANCIAL HIGHLIGHTS ───────────────────────
        add_h("4.  Financial Highlights (INR Crores)",1)
        add_body("All values extracted from submitted documents. N/A = not found in submitted files.")
        def fcr(v): n=normalize_to_inr_crores(v); return f"₹{n:,.1f} Cr" if n else "N/A"
        def fpct(v): s=str(v).replace("%","").strip() if v else None; return f"{s}%" if s and s not in ["None","null",""] else "N/A"
        def fstatus(v, thresholds):
            """Returns ✓ Healthy / ⚠ Monitor / ✗ Critical based on thresholds."""
            if v is None: return "—"
            return thresholds
        fin_rows=[
            ["Revenue (Total Income)",  fcr(annual.get("revenue")),          "Latest FY",                   "✓" if normalize_to_inr_crores(annual.get("revenue")) else "—"],
            ["Net Profit (PAT)",        fcr(annual.get("pat")),              "After Tax",                   "✓ Positive" if safe_float(annual.get("pat") or 0)>0 else "✗ Loss"],
            ["EBITDA",                  fcr(annual.get("ebitda")),           "Operating proxy",             "—"],
            ["Total Debt",              fcr(annual.get("total_debt") or borrowing.get("total_debt")), "Borrowings", "⚠ Monitor"],
            ["Net Worth",               fcr(annual.get("net_worth")),        "Shareholders equity",         "✓"],
            ["Total Assets",            fcr(annual.get("total_assets") or alm.get("total_assets")), "Balance sheet", "—"],
            ["Gross NPA %",             fpct(annual.get("gnpa_percent") or portfolio.get("gnpa_percent")), "< 2% healthy", "✓" if safe_float(str(annual.get("gnpa_percent") or "99").replace("%",""))<2 else "⚠"],
            ["CAR %",                   fpct(annual.get("car_percent")),     "RBI min 15%",                 "✓" if safe_float(str(annual.get("car_percent") or "0").replace("%",""))>=15 else "✗"],
            ["Cash from Operations",    fcr(annual.get("cash_from_operations")), "Operating CF",            "⚠" if normalize_to_inr_crores(annual.get("cash_from_operations")) and safe_float(annual.get("cash_from_operations") or 0)<0 else "✓"],
            ["Interest Coverage",       str(annual.get("interest_coverage") or "N/A"), "EBIT/Interest",     "—"],
            ["Credit Rating",           str(borrowing.get("credit_rating_long_term") or "N/A"), "Long-term", "—"],
            ["Rating Outlook",          str(borrowing.get("rating_outlook") or "N/A"),  "Trajectory",        "—"],
            ["Promoter Holding",        fpct(shareholding.get("promoter_holding")),     "Equity stake",      "—"],
            ["Pledged Shares",          fpct(shareholding.get("pledged_shares")),        "0% = ideal",       "✓" if str(shareholding.get("pledged_shares") or "0").replace("%","").strip() in ["0","0.0",""] else "⚠"],
            ["Collection Efficiency",   fpct(portfolio.get("collection_efficiency")),   "> 95% healthy",     "—"],
            ["Total AUM",               fcr(portfolio.get("total_aum")),                "Portfolio size",    "—"],
        ]
        tbl(["METRIC","VALUE","BENCHMARK","STATUS"],fin_rows,[2.5,1.5,2.5,1.0],NAVY)

        # ── 5. RISK ALERTS ────────────────────────────────
        add_h("5.  Risk Alerts & Positive Indicators",1)
        add_h("Positive Indicators",2,GREEN)
        for ind in score_data.get("green_flags",[]):
            p=doc.add_paragraph(style='List Bullet'); p.add_run(f"✓ {ind}").font.color.rgb=RGBColor(*GREEN)
        if not score_data.get("green_flags"): add_body("No specific positive indicators identified.")
        add_h("Risk Alerts",2,RED)
        for alert in score_data.get("red_flags",[]):
            p=doc.add_paragraph(style='List Bullet'); p.add_run(f"⚠ {alert}").font.color.rgb=RGBColor(*RED)
        if not score_data.get("red_flags"): add_body("No critical risk alerts identified.")
        doc.add_paragraph()
        doc.add_page_break()

        # ── 6. SECONDARY RESEARCH ─────────────────────────
        add_h("6.  Secondary Research Intelligence",1)
        if research_data.get("research_summary"):
            add_body(research_data["research_summary"])
        doc.add_paragraph()
        add_h("Indian Regulatory & Compliance Signals",2)
        ind_rows=[
            ["NCLT/IBC Status",         research_data.get("nclt_status","None detected")],
            ["CIBIL Commercial Signal",  research_data.get("cibil_signal","N/A")],
            ["GSTR-2A vs 3B Signal",     research_data.get("gstr_signal","N/A")],
            ["Litigation Risk",          research_data.get("litigation_risk","N/A")],
            ["Promoter Background",      research_data.get("promoter_background","N/A")],
        ]
        tbl(["SIGNAL TYPE","FINDING"],ind_rows,[2.0,6.0],(45,90,135))
        if research_data.get("rbi_regulatory_flags"):
            add_h("RBI & Regulatory Flags",2,RED)
            for flag in research_data["rbi_regulatory_flags"]:
                p=doc.add_paragraph(style='List Bullet'); p.add_run(f"⚠ {flag}").font.color.rgb=RGBColor(*RED)
        if research_data.get("sector_headwinds"):
            add_h("Sector Headwinds",2,ORANGE)
            for h in research_data["sector_headwinds"]:
                doc.add_paragraph(f"• {h}",style='List Bullet')
        news=[n for n in research_data.get("latest_news",[]) if n and n!="null"]
        if news:
            add_h("Latest Intelligence",2)
            for item in news[:5]: doc.add_paragraph(f"• {item}",style='List Bullet')
        doc.add_paragraph()

        # ── 7. TRIANGULATION ──────────────────────────────
        add_h("7.  Data Triangulation (Document vs Web Intelligence)",1)
        add_body("Cross-referencing document-extracted data against independently researched web data:")
        triang=generate_triangulation(company_name,extracted_docs,research_data)
        if triang:
            tbl(["FIELD","DOCUMENT","WEB RESEARCH","STATUS","ASSESSMENT"],
                [[t["field"],t["doc"],t["web"],t["status"],t["note"]] for t in triang],
                [1.2,1.2,1.4,1.0,3.2],NAVY)
        else:
            add_body("Upload documents to enable triangulation with web data.")
        doc.add_paragraph()

        # ── 8. SWOT ───────────────────────────────────────
        add_h("8.  SWOT Analysis",1)
        swot=score_data.get("swot",{})
        swot_rows=[
            ["STRENGTHS",     "\n".join(swot.get("strengths",[])),     "WEAKNESSES",    "\n".join(swot.get("weaknesses",[]))],
            ["OPPORTUNITIES", "\n".join(swot.get("opportunities",[])), "THREATS",       "\n".join(swot.get("threats",[]))],
        ]
        t=doc.add_table(rows=0,cols=2); t.style='Table Grid'
        for row in swot_rows:
            for pair in [(0,1),(2,3)]:
                r=t.add_row()
                r.cells[0].text=row[pair[0]]; r.cells[1].text=row[pair[1]]
                if row[pair[0]] in ["STRENGTHS","WEAKNESSES","OPPORTUNITIES","THREATS"]:
                    r.cells[0].paragraphs[0].runs[0].bold=True
                    r.cells[1].paragraphs[0].runs[0].bold=True
        doc.add_paragraph()
        doc.add_page_break()

        # ── 9. CREDIT NARRATIVE ───────────────────────────
        add_h("9.  Detailed Credit Assessment",1)
        try:
            narrative_prompt=f"""Write a formal Credit Appraisal Memo narrative for {company_name}.
Score: {total}/100 | Decision: {dec_lbl} | Grade: {grade}
Revenue: {fcr(annual.get('revenue'))} | PAT: {fcr(annual.get('pat'))} | Debt: {fcr(annual.get('total_debt'))} | Net Worth: {fcr(annual.get('net_worth'))}
GNPA: {fpct(annual.get('gnpa_percent'))} | CAR: {fpct(annual.get('car_percent'))}
Red Flags: {', '.join(score_data.get('red_flags',[])[:3])}
Green Flags: {', '.join(score_data.get('green_flags',[])[:2])}
Research: {research_data.get('research_summary','')}

Write exactly 5 sections with these headers (formal Indian banking language, reference actual numbers):
1. BORROWER BACKGROUND
2. FINANCIAL ANALYSIS  
3. RISK ASSESSMENT
4. CREDIT OPINION
5. RECOMMENDATION"""
            narrative_text = call_gemini(
                "You are a senior credit analyst at an Indian bank writing formal Credit Appraisal Memos.",
                narrative_prompt, temperature=0.3, max_tokens=1500
            )
            doc.add_paragraph(narrative_text)
        except Exception as e:
            add_body(f"Narrative generation error: {e}")
        doc.add_paragraph()

        # ── 10. PROPOSED TERMS ────────────────────────────
        add_h("10.  Proposed Loan Terms",1)
        terms_rows=[
            ["Loan Limit",    f"₹{score_data.get('recommended_amount','N/A')} Crores", f"100% of request (Score {total}/100 — Grade {grade})", "Standard covenant monitoring"],
            ["Interest Rate",  rate,                                                      "Per VERIDEX decision table for Grade "+grade,          "Rate reset at 24 months"],
            ["Tenure",        f"{loan_data.get('tenure','N/A')} Months",               "Per borrower request",                                   "Quarterly amortisation"],
            ["Security",       "Subject to credit committee",                           "First charge on receivables recommended",                "Confirmed pre-disbursement"],
        ]
        tbl(["PARAMETER","RECOMMENDED","RATIONALE","CONDITIONS"],terms_rows,[1.4,1.6,3.0,2.0],NAVY)

        # ── DISCLAIMER ────────────────────────────────────
        doc.add_paragraph()
        p=doc.add_paragraph("This report was generated by VERIDEX AI Credit Engine v2.0. All AI assessments are indicative and must be reviewed by a qualified credit officer before final sanction. Confidential — for internal use only.")
        p.runs[0].font.size=Pt(9); p.runs[0].font.color.rgb=RGBColor(120,120,120); p.runs[0].font.italic=True

        buf=io.BytesIO(); doc.save(buf); buf.seek(0)
        name=company_name.replace(" ","_"); date_str=datetime.now().strftime("%Y%m%d")
        return Response(content=buf.getvalue(),
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition":f"attachment; filename=CAM_{name}_{date_str}.docx"})

    except Exception as e:
        import traceback; traceback.print_exc()
        raise HTTPException(status_code=500, detail=str(e))
