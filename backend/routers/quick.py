from fastapi import APIRouter, UploadFile, File, Form, HTTPException, Response
import uuid, os, json, re, asyncio, io
from typing import Dict
from datetime import datetime

try:
    import pytesseract; OCR_AVAILABLE = True
except ImportError: OCR_AVAILABLE = False

router = APIRouter(prefix="/api/quick")
quick_session = {}

def robust_json_parser(text):
    if not text: return {}
    clean = text.replace("```json","").replace("```","").strip()
    try: return json.loads(clean)
    except: pass
    try:
        m = re.search(r'\{.*\}', clean, re.DOTALL)
        if m: return json.loads(m.group())
    except: pass
    return {}

def extract_text_from_file(contents: bytes, filename: str) -> str:
    MAX_PAGES = 20
    if filename.lower().endswith('.pdf'):
        import pdfplumber
        text = ""
        with pdfplumber.open(io.BytesIO(contents)) as pdf:
            total = len(pdf.pages)
            # Smart sample: first 12 + last 8 for large docs
            if total <= MAX_PAGES:
                pages = pdf.pages[:total]
            else:
                first = pdf.pages[:12]
                last  = pdf.pages[max(12, total-8):]
                pages = list(first) + list(last)
            for page in pages:
                t = page.extract_text() or ""
                if len(t.strip()) > 30: text += t + "\n"
                elif OCR_AVAILABLE:
                    try:
                        img = page.to_image(resolution=150).original
                        import pytesseract
                        text += pytesseract.image_to_string(img) + "\n"
                    except: pass
        return text.strip()[:25000]
    elif filename.lower().endswith('.docx'):
        from docx import Document as DocxDoc
        d = DocxDoc(io.BytesIO(contents))
        return "\n".join([p.text for p in d.paragraphs])[:25000]
    elif filename.lower().endswith(('.xlsx','.xls')):
        import pandas as pd
        dfs = pd.read_excel(io.BytesIO(contents), sheet_name=None)
        return "\n\n".join([f"Sheet: {sh}\n{df.to_string()}" for sh, df in dfs.items()])[:25000]
    try: return contents.decode('utf-8')[:25000]
    except: return ""

def extract_any_document(text: str, company_name: str, gemini_caller) -> dict:
    prompt = f"""You are a senior Indian credit analyst. Extract ALL financial data from this {company_name} document.

CRITICAL UNIT RULES:
- All values MUST be in INR Crores.
- USD billions × 8300 = INR Cr. USD millions × 83 / 100 = INR Cr.
- INR Lakhs ÷ 100 = INR Cr.
- If you are unsure of units, use null — never guess.

Return ONLY valid JSON:
{{
  "annual_report": {{
    "revenue": null, "pat": null, "pbt": null, "ebitda": null, "total_debt": null,
    "net_worth": null, "total_assets": null, "gnpa_percent": null, "car_percent": null,
    "interest_coverage": null, "cash_from_operations": null, "auditor_remarks": null
  }},
  "borrowing_profile": {{
    "total_debt": null, "credit_rating_long_term": null, "rating_outlook": null,
    "average_cost_of_funds": null
  }},
  "shareholding_pattern": {{
    "promoter_holding": null, "pledged_shares": null, "fii_holding": null
  }},
  "portfolio_cuts": {{
    "gnpa_percent": null, "nnpa_percent": null, "collection_efficiency": null, "total_aum": null
  }},
  "alm_statement": {{
    "total_assets": null, "total_liabilities": null, "liquidity_gap": null
  }}
}}

Document Text:
{text[:20000]}"""

    response_text = gemini_caller(
        "You are a senior Indian credit analyst extracting financial data. Return ONLY valid JSON.",
        prompt, temperature=0.05, max_tokens=2500
    )
    return robust_json_parser(response_text)

@router.post("/upload")
async def quick_upload(file: UploadFile = File(...), session_id: str = Form(...)):
    contents = await file.read()
    text = extract_text_from_file(contents, file.filename)
    quick_session[session_id] = {"text": text, "filename": file.filename}
    return {"filename": file.filename, "characters_extracted": len(text), "preview": text[:300]}

@router.post("/analyze")
async def quick_analyze(company_name: str = Form(...), session_id: str = Form(...)):
    if session_id not in quick_session:
        raise HTTPException(status_code=404, detail="Session not found")
    from .report import call_gemini, normalize_to_inr_crores

    text = quick_session[session_id]["text"]
    extracted = extract_any_document(text, company_name, call_gemini)
    quick_session[session_id]["extracted"] = extracted
    quick_session[session_id]["company_name"] = company_name

    annual      = extracted.get("annual_report", {})
    borrowing   = extracted.get("borrowing_profile", {})
    portfolio   = extracted.get("portfolio_cuts", {})
    alm         = extracted.get("alm_statement", {})
    shareholding= extracted.get("shareholding_pattern", {})

    def norm(v): return normalize_to_inr_crores(v)
    def pct(v):
        sv = str(v).replace("%","").strip() if v else None
        return f"{sv}%" if sv and sv not in ["None","null",""] else None

    rev  = norm(annual.get("revenue"))
    pat  = norm(annual.get("pat"))
    debt = norm(annual.get("total_debt") or borrowing.get("total_debt"))
    nw   = norm(annual.get("net_worth"))
    assets = norm(annual.get("total_assets") or alm.get("total_assets"))

    # Compute ratios from normalized values
    de_ratio = round(float(debt)/float(nw), 4) if debt and nw and float(nw) > 0 else None
    roe      = round((float(pat)/float(nw))*100, 2) if pat and nw and float(nw) > 0 else None

    financials = {}
    if rev   is not None: financials["Revenue"]      = f"₹{rev} Cr"
    if pat   is not None: financials["Profit (PAT)"] = f"₹{pat} Cr"
    if annual.get("ebitda"):   financials["EBITDA"]  = f"₹{norm(annual.get('ebitda'))} Cr"
    if debt  is not None: financials["Total Debt"]   = f"₹{debt} Cr"
    if nw    is not None: financials["Net Worth"]    = f"₹{nw} Cr"
    if assets is not None: financials["Total Assets"]= f"₹{assets} Cr"
    if de_ratio is not None: financials["D/E Ratio"] = f"{de_ratio}x"
    if roe      is not None: financials["ROE"]        = f"{roe}%"
    if annual.get("gnpa_percent"): financials["GNPA %"]= pct(annual.get("gnpa_percent"))
    if annual.get("car_percent"):  financials["CAR %"] = pct(annual.get("car_percent"))
    if borrowing.get("credit_rating_long_term"): financials["Credit Rating"] = borrowing["credit_rating_long_term"]
    if borrowing.get("rating_outlook"):          financials["Rating Outlook"] = borrowing["rating_outlook"]
    if shareholding.get("promoter_holding"):     financials["Promoter Holding"] = pct(shareholding["promoter_holding"])

    financials = {k: v for k, v in financials.items() if v is not None}
    quick_session[session_id]["financials"] = financials
    return {"company": company_name, "financials": financials}

@router.post("/score")
async def quick_score(
    session_id: str = Form(...),
    loan_amount: str = Form("50"),
    tenure: str = Form("36"),
    interest_rate: str = Form("11.5"),
    sector: str = Form("NBFC")
):
    if session_id not in quick_session:
        raise HTTPException(status_code=404, detail="Session not found")
    from .report import calculate_universal_score, generate_swot

    session  = quick_session[session_id]
    extracted = session.get("extracted", {})
    company   = session.get("company_name", "Unknown Company")

    entity_data = {
        "company_name": company, "sector": sector,
        "loan_amount": float(loan_amount),
        "interest_rate": float(interest_rate),
        "tenure": int(tenure)
    }
    scoring = calculate_universal_score(company, extracted, [], entity_data)
    scoring["swot"] = generate_swot(company, extracted, [])
    quick_session[session_id]["scoring"] = scoring
    quick_session[session_id]["entity_data"] = entity_data
    return scoring

@router.post("/research")
async def quick_research(session_id: str = Form(...)):
    if session_id not in quick_session:
        raise HTTPException(status_code=404, detail="Session not found")
    from .report import cached_tavily, call_gemini, get_decision, get_risk_level

    session = quick_session[session_id]
    company = session.get("company_name")

    loop = asyncio.get_event_loop()
    q1 = f"{company} India fraud litigation NCLT default NPA 2024 2025"
    q2 = f"{company} India financial news credit rating revenue profit 2024 2025"
    q3 = f"{company} India RBI SEBI regulatory action penalty 2024"
    q4 = f"{company} India promoter background news"

    try:
        results = await asyncio.wait_for(asyncio.gather(
            loop.run_in_executor(None, lambda: cached_tavily(q1, 3)),
            loop.run_in_executor(None, lambda: cached_tavily(q2, 3)),
        ), timeout=20.0)
    except asyncio.TimeoutError:
        results = [{'results':[]}, {'results':[]}]

    all_results = []
    seen = set()
    for res in results:
        for r in res.get("results", []):
            if r.get("url") not in seen:
                all_results.append(r); seen.add(r.get("url"))

    findings_text = "\n".join([f"- {r.get('title')}: {r.get('content','')[:300]}" for r in all_results])

    prompt = f"""You are a credit risk analyst. Based on web results about {company}:
1. RISK LEVEL: (LOW / MEDIUM / HIGH / CRITICAL)
2. KEY FINDINGS: max 5 bullet points (use * prefix)
3. SUMMARY: 2-3 sentence overall assessment

Results: {findings_text[:5000]}
Be factual. Only report what is in the results."""

    text = call_gemini(
        "You are a credit risk analyst. Analyze web search results and provide a structured risk assessment.",
        prompt, temperature=0.3, max_tokens=1000
    )
    risk_level = "MEDIUM"
    if "RISK LEVEL: LOW"      in text.upper(): risk_level = "LOW"
    elif "RISK LEVEL: MEDIUM" in text.upper(): risk_level = "MEDIUM"
    elif "RISK LEVEL: HIGH"   in text.upper(): risk_level = "HIGH"
    elif "RISK LEVEL: CRITICAL" in text.upper(): risk_level = "CRITICAL"

    # Adjust score based on risk
    scoring = session.get("scoring", {})
    adj = {"CRITICAL": -20, "HIGH": -10, "MEDIUM": 0, "LOW": 5}.get(risk_level, 0)
    if adj != 0:
        new_score = max(0, min(100, scoring.get("score", 0) + adj))
        from .report import get_decision, get_risk_level as grl
        decision, grade, rate = get_decision(new_score)
        scoring.update({"score": new_score, "decision": decision, "grade": grade,
                        "recommended_rate": rate, "risk_level": grl(new_score)})
        quick_session[session_id]["scoring"] = scoring

    report_id = str(uuid.uuid4())[:8]
    quick_session[session_id]["research"] = {
        "risk_level": risk_level, "summary": text,
        "sources": len(all_results), "report_id": report_id,
        "findings": all_results
    }
    return {"risk_level": risk_level, "summary": text, "sources": len(all_results), "report_id": report_id}

@router.get("/report/{report_id}")
async def get_quick_report(report_id: str):
    from .report import generate_swot, normalize_to_inr_crores, get_decision
    from docx.shared import RGBColor, Inches, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    session_found = next(
        (d for d in quick_session.values() if d.get("research",{}).get("report_id") == report_id),
        None
    )
    if not session_found:
        raise HTTPException(status_code=404, detail="Report not found")

    data       = session_found
    company    = data.get("company_name","Unknown")
    scoring    = data.get("scoring", {})
    extracted  = data.get("extracted", {})
    research   = data.get("research", {})
    findings   = research.get("findings", [])
    entity_data= data.get("entity_data", {})
    financials = data.get("financials", {})

    doc = Document()
    section = doc.sections[0]
    section.left_margin = Inches(1); section.right_margin = Inches(1)

    def add_h(text, level=1, color=(26,58,107)):
        p = doc.add_heading(text, level=level)
        for run in p.runs: run.font.color.rgb = RGBColor(*color)

    def add_kv(label, value):
        p = doc.add_paragraph()
        r1 = p.add_run(f"{label}: "); r1.bold = True; r1.font.size = Pt(11)
        p.add_run(str(value) if value is not None else "N/A").font.size = Pt(11)

    title = doc.add_heading("Quick Credit Appraisal Memo", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub = doc.add_paragraph("VERIDEX® — CONFIDENTIAL")
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()
    add_kv("Entity", company)
    add_kv("Sector", entity_data.get("sector","N/A"))
    add_kv("Loan Amount", f"₹{entity_data.get('loan_amount','N/A')} Cr")
    add_kv("Report Date", datetime.now().strftime("%d %B %Y"))
    add_kv("Generated By", "VERIDEX AI Quick Appraisal v2.0")
    doc.add_page_break()

    add_h("1. Credit Decision", 1)
    add_kv("Decision", scoring.get("decision","N/A"))
    add_kv("Grade",    scoring.get("grade","N/A"))
    add_kv("Intelli-Score", f"{scoring.get('score',0)}/100")
    add_kv("Risk Level", scoring.get("risk_level","N/A"))
    add_kv("Recommended Limit", f"₹{scoring.get('recommended_amount','N/A')} Cr")
    add_kv("Recommended Rate",  scoring.get("recommended_rate","N/A"))
    doc.add_paragraph()
    p = doc.add_paragraph(); p.add_run("Reasoning: ").bold = True; p.add_run(scoring.get("reasoning",""))

    add_h("2. Financial Summary", 1)
    table = doc.add_table(rows=len(financials)+1, cols=2); table.style = "Table Grid"
    hdr = table.rows[0].cells; hdr[0].text = "Metric"; hdr[1].text = "Value"
    for i, (k, v) in enumerate(financials.items()):
        row = table.rows[i+1].cells; row[0].text = k; row[1].text = str(v)
    doc.add_paragraph()

    add_h("3. Five Cs Analysis", 1)
    five_cs = scoring.get("five_cs", {})
    for c, val in five_cs.items():
        if isinstance(val, dict):
            add_kv(c.capitalize(), f"{val.get('score',0)}/{val.get('max',20)} pts — {val.get('notes','')}")
        else:
            add_kv(c.capitalize(), f"{val} pts")

    add_h("4. Risk Alerts", 1)
    for flag in scoring.get("red_flags",[]):
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(f"⚠ {flag}").font.color.rgb = RGBColor(180,0,0)
    add_h("Positive Indicators", 2)
    for flag in scoring.get("green_flags",[]):
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(f"✓ {flag}").font.color.rgb = RGBColor(0,128,0)

    swot = scoring.get("swot", {})
    if swot:
        add_h("5. SWOT Analysis", 1)
        for quad, items in swot.items():
            add_kv(quad.upper(), ", ".join(items) if items else "N/A")

    add_h("6. Web Intelligence", 1)
    add_kv("Risk Level", research.get("risk_level","N/A"))
    doc.add_paragraph(research.get("summary",""))
    if findings:
        add_h("Research Sources", 2)
        for f in findings[:5]:
            p = doc.add_paragraph(); p.add_run(f.get("title","")).bold = True
            doc.add_paragraph((f.get("content") or f.get("snippet") or "")[:300])

    buffer = io.BytesIO(); doc.save(buffer); buffer.seek(0)
    return Response(
        content=buffer.getvalue(),
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f"attachment; filename=QuickCAM_{company.replace(' ','_')}.docx"}
    )
